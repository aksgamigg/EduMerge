# =============================================================================
# IMPORTS
# =============================================================================
from tkinter import filedialog, messagebox, colorchooser
from tkinter import Text, END, LEFT, RIGHT, BOTTOM, TOP, X, Y, BOTH, W, E
import ttkbootstrap as ttk
from ttkbootstrap import Toplevel, Label, PhotoImage
from ttkbootstrap.dialogs import Messagebox
from ttkbootstrap.widgets.tableview import Tableview        # fixed deprecated import
from ttkbootstrap.constants import *
from pathlib import Path
import os, platform, subprocess
from PIL import Image, ImageTk
import csv
from docx import Document
from docx.shared import Pt, RGBColor
import pypdf


# =============================================================================
# Text Editor
# =============================================================================
class ModernEditorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("EduText: Text Editor")
        self.root.geometry("1200x800")
        try:
            self.root.iconbitmap("logo.ico")
        except:
            pass

        self.filename = None
        self.is_saved = True
        self.images = []
        self.current_file_type = "text"

        # Store CSV data
        self.csv_data = []
        self.csv_headers = []

        # Apply custom styling (safe)
        self.apply_custom_styles()

        # Create modern UI
        self.create_modern_sidebar()
        self.create_main_area()
        self.create_status_bar()

        # Bind shortcuts
        self.root.bind("<Control-s>", lambda e: self.save_file())
        self.root.bind("<Control-o>", lambda e: self.open_file())
        self.root.bind("<Control-n>", lambda e: self.new_file())
        self.root.bind("<Control-b>", lambda e: self.apply_bold())
        self.root.bind("<Control-i>", lambda e: self.apply_italic())
        self.root.bind("<Control-u>", lambda e: self.apply_underline())

        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def apply_custom_styles(self):
        """Apply safer custom styles without using unsupported 'prefix' builder names."""
        style = ttk.Style()

        # Configure base TButton padding so buttons look a bit larger / rounded by theme
        # Note: avoid leading custom prefixes like "Rounded." which trigger missing builder methods.
        style.configure('TButton', padding=(12, 8), borderwidth=0, relief="flat")

        # If you want variant-specific looks, use bootstyle values at widget level (e.g. bootstyle="success")
        # or create fully named styles that match standard ttk patterns (not custom prefixes that expect builders).

        # Example: small button style (named safely)
        style.configure('Small.TButton', padding=(8, 6), borderwidth=0)

        # You can still use bootstyle param when creating widgets, e.g. bootstyle="success" or "primary-outline".

    def create_modern_sidebar(self):
        # Sidebar with gradient-like effect
        self.sidebar = ttk.Frame(self.root, bootstyle="dark", width=280)
        self.sidebar.pack(side=LEFT, fill=Y, padx=0, pady=0)
        self.sidebar.pack_propagate(False)

        # Logo/Title area with modern card design
        title_frame = ttk.Frame(self.sidebar, bootstyle="dark")
        title_frame.pack(fill=X, padx=25, pady=35)

        # App icon and title
        ttk.Label(
            title_frame,
            text="📝",
            font=("Segoe UI Emoji", 32),
            bootstyle="light"
        ).pack()

        ttk.Label(
            title_frame,
            text="EduText",
            font=("Segoe UI", 22, "bold"),
            bootstyle="light",
            foreground="#FFFFFF"
        ).pack(pady=(8, 2))

        ttk.Label(
            title_frame,
            text="A Student Software",
            font=("Segoe UI", 9),
            bootstyle="light",
            foreground="#A0A0A0"
        ).pack()

        # Separator
        ttk.Separator(self.sidebar, bootstyle="secondary").pack(fill=X, padx=25, pady=15)

        # File operations section
        self.create_section_header("FILE OPERATIONS")
        self.create_sidebar_button("📄  New Document", self.new_file, "success")
        self.create_sidebar_button("📂  Open File", self.open_file, "info")
        self.create_sidebar_button("💾  Save", self.save_file, "primary")
        self.create_sidebar_button("💾  Save As", self.save_as_file, "primary")

        # Insert section
        self.create_section_header("INSERT CONTENT")
        self.create_sidebar_button("🖼️  Insert Image", self.insert_image, "warning")
        self.create_sidebar_button("📊  CSV Table", self.create_csv_table, "info")

        # Export section
        self.create_section_header("EXPORT OPTIONS")
        self.create_sidebar_button("📄  Export DOCX", self.export_to_docx, "danger")
        self.create_sidebar_button("📑  Export PDF", self.export_to_pdf, "danger")

        # Footer with version
        footer = ttk.Frame(self.sidebar, bootstyle="dark")
        footer.pack(side=BOTTOM, fill=X, padx=25, pady=20)

        ttk.Label(
            footer,
            text="v0.2 • Modern Edition",
            font=("Segoe UI", 8),
            bootstyle="light",
            foreground="#707070"
        ).pack()

    def create_section_header(self, text):
        """Create a modern section header"""
        header_frame = ttk.Frame(self.sidebar, bootstyle="dark")
        header_frame.pack(fill=X, padx=25, pady=(20, 12))

        ttk.Label(
            header_frame,
            text=text,
            font=("Segoe UI", 8, "bold"),
            bootstyle="light",
            foreground="#808080"
        ).pack(anchor=W)

    def create_sidebar_button(self, text, command, style):
        """Create modern rounded sidebar buttons using built-in bootstyles (no custom prefix)."""
        # Use bootstyle directly (e.g. "success", "info", "primary") rather than "Rounded.success"
        btn = ttk.Button(
            self.sidebar,
            text=text,
            command=command,
            bootstyle=style,
            width=28,
            cursor="hand2"
        )
        btn.pack(padx=25, pady=4, fill=X)

    def create_main_area(self):
        # Main container with padding
        main_container = ttk.Frame(self.root, bootstyle="light")
        main_container.pack(side=LEFT, fill=BOTH, expand=True)

        # Modern toolbar with card design
        self.create_modern_toolbar(main_container)

        # Content area with modern card wrapper
        content_wrapper = ttk.Frame(main_container, bootstyle="light")
        content_wrapper.pack(fill=BOTH, expand=True, padx=25, pady=(15, 20))

        # Modern card container
        card_frame = ttk.Frame(content_wrapper, bootstyle="light", relief="flat")
        card_frame.pack(fill=BOTH, expand=True)

        # Notebook for different views
        self.notebook = ttk.Notebook(card_frame, bootstyle="dark")
        self.notebook.pack(fill=BOTH, expand=True)

        # Text editor tab
        self.text_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.text_frame, text="✏️  Text Editor")

        self.text_area = Text(
            self.text_frame,
            wrap="word",
            font=("Segoe UI", 12),
            bg="#FFFFFF",
            fg="#1A1A1A",
            insertbackground="#3498DB",
            relief="flat",
            padx=30,
            pady=30,
            selectbackground="#3498DB",
            selectforeground="white",
            spacing1=3,
            spacing3=3
        )

        text_scrollbar = ttk.Scrollbar(self.text_frame, command=self.text_area.yview, bootstyle="round")
        self.text_area.configure(yscrollcommand=text_scrollbar.set)

        text_scrollbar.pack(side=RIGHT, fill=Y)
        self.text_area.pack(fill=BOTH, expand=True)

        # Configure text tags with modern styling
        self.text_area.tag_configure("bold", font=("Segoe UI", 12, "bold"))
        self.text_area.tag_configure("italic", font=("Segoe UI", 12, "italic"))
        self.text_area.tag_configure("underline", font=("Segoe UI", 12, "underline"))
        self.text_area.tag_configure("heading1", font=("Segoe UI", 32, "bold"), foreground="#E74C3C", spacing3=15)
        self.text_area.tag_configure("heading2", font=("Segoe UI", 24, "bold"), foreground="#3498DB", spacing3=12)
        self.text_area.tag_configure("heading3", font=("Segoe UI", 18, "bold"), foreground="#16A085", spacing3=10)

        self.text_area.bind("<<Modified>>", self.on_text_change)
        self.text_area.bind("<KeyRelease>", self.update_status)

        # CSV table tab
        self.csv_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.csv_frame, text="📊  CSV Table")

        # PDF viewer tab
        self.pdf_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.pdf_frame, text="📑  PDF Viewer")

        self.pdf_text = Text(
            self.pdf_frame,
            wrap="word",
            font=("Segoe UI", 11),
            bg="#FAFAFA",
            fg="#1A1A1A",
            relief="flat",
            padx=30,
            pady=30,
            state="disabled"
        )
        pdf_scrollbar = ttk.Scrollbar(self.pdf_frame, command=self.pdf_text.yview, bootstyle="round")
        self.pdf_text.configure(yscrollcommand=pdf_scrollbar.set)
        pdf_scrollbar.pack(side=RIGHT, fill=Y)
        self.pdf_text.pack(fill=BOTH, expand=True)

    def create_modern_toolbar(self, parent):
        """Create modern toolbar with rounded buttons"""
        toolbar_container = ttk.Frame(parent, bootstyle="light")
        toolbar_container.pack(fill=X, padx=25, pady=(25, 5))

        toolbar = ttk.Frame(toolbar_container, bootstyle="light", relief="flat")
        toolbar.pack(fill=X, pady=5)

        # Font Controls Card
        font_card = ttk.Labelframe(
            toolbar,
            text="  Font  ",
            bootstyle="primary",
            padding=15
        )
        font_card.pack(side=LEFT, padx=(0, 10))

        # Font family with modern styling
        self.font_family_var = ttk.StringVar(value="Segoe UI")
        fonts = ["Segoe UI", "Arial", "Times New Roman", "Calibri", "Georgia", "Courier New", "Consolas"]
        font_combo = ttk.Combobox(
            font_card,
            textvariable=self.font_family_var,
            values=fonts,
            width=14,
            state="readonly",
            bootstyle="primary"
        )
        font_combo.pack(side=LEFT, padx=(0, 8))
        font_combo.bind("<<ComboboxSelected>>", lambda e: self.change_font())

        # Font size
        self.font_size_var = ttk.IntVar(value=12)
        size_spinbox = ttk.Spinbox(
            font_card,
            from_=8,
            to=72,
            textvariable=self.font_size_var,
            width=6,
            command=self.change_font,
            bootstyle="primary"
        )
        size_spinbox.pack(side=LEFT, padx=2)

        # Text Style Card
        style_card = ttk.Labelframe(
            toolbar,
            text="  Style  ",
            bootstyle="success",
            padding=15
        )
        style_card.pack(side=LEFT, padx=5)

        # Rounded style buttons (use bootstyle variants)
        btn_bold = ttk.Button(
            style_card,
            text="𝐁",
            bootstyle="success-outline",
            command=self.apply_bold,
            width=4,
            cursor="hand2"
        )
        btn_bold.pack(side=LEFT, padx=2)

        btn_italic = ttk.Button(
            style_card,
            text="𝐼",
            bootstyle="success-outline",
            command=self.apply_italic,
            width=4,
            cursor="hand2"
        )
        btn_italic.pack(side=LEFT, padx=2)

        btn_underline = ttk.Button(
            style_card,
            text="U̲",
            bootstyle="success-outline",
            command=self.apply_underline,
            width=4,
            cursor="hand2"
        )
        btn_underline.pack(side=LEFT, padx=2)

        # Heading Card
        heading_card = ttk.Labelframe(
            toolbar,
            text="  Headings  ",
            bootstyle="info",
            padding=15
        )
        heading_card.pack(side=LEFT, padx=5)

        ttk.Button(
            heading_card,
            text="H1",
            bootstyle="danger",
            command=lambda: self.apply_heading("heading1"),
            width=4,
            cursor="hand2"
        ).pack(side=LEFT, padx=2)

        ttk.Button(
            heading_card,
            text="H2",
            bootstyle="primary",
            command=lambda: self.apply_heading("heading2"),
            width=4,
            cursor="hand2"
        ).pack(side=LEFT, padx=2)

        ttk.Button(
            heading_card,
            text="H3",
            bootstyle="success",
            command=lambda: self.apply_heading("heading3"),
            width=4,
            cursor="hand2"
        ).pack(side=LEFT, padx=2)

        # Color Card
        color_card = ttk.Labelframe(
            toolbar,
            text="  Colors  ",
            bootstyle="warning",
            padding=15
        )
        color_card.pack(side=LEFT, padx=5)

        ttk.Button(
            color_card,
            text="🎨 Text",
            bootstyle="warning",
            command=self.change_text_color,
            width=9,
            cursor="hand2"
        ).pack(side=LEFT, padx=2)

        ttk.Button(
            color_card,
            text="🎨 BG",
            bootstyle="warning-outline",
            command=self.change_bg_color,
            width=8,
            cursor="hand2"
        ).pack(side=LEFT, padx=2)

    def create_status_bar(self):
        """Create modern status bar"""
        status_frame = ttk.Frame(self.root, bootstyle="dark")
        status_frame.pack(fill=X, side=BOTTOM)

        # Left side - document stats
        self.status_bar = ttk.Label(
            status_frame,
            text="📄 Lines: 1  |  Characters: 0  |  Words: 0  |  Ready",
            anchor=W,
            font=("Segoe UI", 9),
            padding=(15, 8)
        )
        self.status_bar.pack(side=LEFT, fill=X, expand=True)

        # Right side - file type indicator
        self.file_type_label = ttk.Label(
            status_frame,
            text="📝 Text Document",
            anchor=E,
            font=("Segoe UI", 9, "bold"),
            padding=(15, 8)
        )
        self.file_type_label.pack(side=RIGHT)

    def apply_bold(self):
        try:
            current_tags = self.text_area.tag_names("sel.first")
            if "bold" in current_tags:
                self.text_area.tag_remove("bold", "sel.first", "sel.last")
            else:
                self.text_area.tag_add("bold", "sel.first", "sel.last")
        except:
            pass

    def apply_italic(self):
        try:
            current_tags = self.text_area.tag_names("sel.first")
            if "italic" in current_tags:
                self.text_area.tag_remove("italic", "sel.first", "sel.last")
            else:
                self.text_area.tag_add("italic", "sel.first", "sel.last")
        except:
            pass

    def apply_underline(self):
        try:
            current_tags = self.text_area.tag_names("sel.first")
            if "underline" in current_tags:
                self.text_area.tag_remove("underline", "sel.first", "sel.last")
            else:
                self.text_area.tag_add("underline", "sel.first", "sel.last")
        except:
            pass

    def apply_heading(self, heading_tag):
        try:
            for tag in ["heading1", "heading2", "heading3"]:
                self.text_area.tag_remove(tag, "sel.first", "sel.last")
            self.text_area.tag_add(heading_tag, "sel.first", "sel.last")
        except:
            pass

    def change_font(self):
        family = self.font_family_var.get()
        size = self.font_size_var.get()

        self.text_area.configure(font=(family, size))
        self.text_area.tag_configure("bold", font=(family, size, "bold"))
        self.text_area.tag_configure("italic", font=(family, size, "italic"))
        self.text_area.tag_configure("underline", font=(family, size, "underline"))
        self.text_area.tag_configure("heading1", font=(family, size + 20, "bold"))
        self.text_area.tag_configure("heading2", font=(family, size + 12, "bold"))
        self.text_area.tag_configure("heading3", font=(family, size + 6, "bold"))

    def change_text_color(self):
        try:
            color = colorchooser.askcolor(title="Choose Text Color")
            if color[1]:
                tag_name = f"color_{color[1]}"
                self.text_area.tag_configure(tag_name, foreground=color[1])
                self.text_area.tag_add(tag_name, "sel.first", "sel.last")
        except:
            pass

    def change_bg_color(self):
        try:
            color = colorchooser.askcolor(title="Choose Background Color")
            if color[1]:
                tag_name = f"bg_{color[1]}"
                self.text_area.tag_configure(tag_name, background=color[1])
                self.text_area.tag_add(tag_name, "sel.first", "sel.last")
        except:
            pass

    def insert_image(self):
        filepath = filedialog.askopenfilename(
            title="Select Image",
            filetypes=[("Image Files", "*.png *.jpg *.jpeg *.gif *.bmp"), ("All Files", "*.*")]
        )

        if filepath:
            try:
                img = Image.open(filepath)
                max_width = 700
                if img.width > max_width:
                    ratio = max_width / img.width
                    new_height = int(img.height * ratio)
                    # Pillow >= 9.1: Image.Resampling.LANCZOS; fallback to Image.LANCZOS if necessary
                    try:
                        resample = Image.Resampling.LANCZOS
                    except AttributeError:
                        resample = Image.LANCZOS
                    img = img.resize((max_width, new_height), resample)

                photo = ImageTk.PhotoImage(img)
                self.images.append(photo)
                self.text_area.image_create("insert", image=photo)
                self.text_area.insert("insert", "\n")
                self.is_saved = False
                self.update_title()
            except Exception as e:
                messagebox.showerror("Error", f"Could not insert image:\n{e}")

    def create_csv_table(self):
        """Switch to CSV view and create empty table"""
        self.notebook.select(1)
        self.current_file_type = "csv"
        self.file_type_label.config(text="📊 CSV Table")

    def open_csv(self, filepath):
        """Open and display CSV file"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                self.csv_data = list(csv_reader)

                if self.csv_data:
                    self.csv_headers = self.csv_data[0]

                    # Clear previous table
                    for widget in self.csv_frame.winfo_children():
                        widget.destroy()

                    # Create tableview
                    coldata = [{"text": header, "stretch": True} for header in self.csv_headers]
                    rowdata = self.csv_data[1:]

                    table = Tableview(
                        self.csv_frame,
                        coldata=coldata,
                        rowdata=rowdata,
                        paginated=True,
                        searchable=True,
                        bootstyle="primary",
                        height=20
                    )
                    table.pack(fill=BOTH, expand=True, padx=20, pady=20)

                    self.notebook.select(1)
                    self.current_file_type = "csv"
                    self.file_type_label.config(text="📊 CSV Table")
        except Exception as e:
            messagebox.showerror("Error", f"Could not open CSV:\n{e}")

    def open_docx(self, filepath):
        """Open and display DOCX file"""
        try:
            doc = Document(filepath)
            self.text_area.delete("1.0", END)

            for para in doc.paragraphs:
                self.text_area.insert(END, para.text + "\n")

            self.notebook.select(0)
            self.current_file_type = "docx"
            self.file_type_label.config(text="📄 Word Document")
        except Exception as e:
            messagebox.showerror("Error", f"Could not open DOCX:\n{e}")

    def open_pdf(self, filepath):
        """Open and display PDF file"""
        try:
            reader = pypdf.PdfReader(filepath)
            self.pdf_text.config(state="normal")
            self.pdf_text.delete("1.0", END)

            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                self.pdf_text.insert(END, f"━━━ Page {page_num} ━━━\n\n")
                self.pdf_text.insert(END, text + "\n\n")

            self.pdf_text.config(state="disabled")
            self.notebook.select(2)
            self.current_file_type = "pdf"
            self.file_type_label.config(text="📑 PDF Document")
        except Exception as e:
            messagebox.showerror("Error", f"Could not open PDF:\n{e}")

    def new_file(self):
        if not self.is_saved:
            response = messagebox.askyesnocancel("Save Changes", "Do you want to save changes?")
            if response is None:
                return
            elif response:
                self.save_file()

        self.text_area.delete("1.0", END)
        self.images.clear()
        self.csv_data.clear()
        self.filename = None
        self.is_saved = True
        self.current_file_type = "text"
        self.file_type_label.config(text="📝 Text Document")
        self.notebook.select(0)
        self.update_status()

    def open_file(self):
        filepath = filedialog.askopenfilename(
            defaultextension=".txt",
            filetypes=[
                ("All Supported", "*.txt *.csv *.docx *.pdf"),
                ("Text Files", "*.txt"),
                ("CSV Files", "*.csv"),
                ("Word Documents", "*.docx"),
                ("PDF Files", "*.pdf"),
                ("All Files", "*.*")
            ]
        )

        if filepath:
            self.filename = filepath
            ext = os.path.splitext(filepath)[1].lower()

            if ext == '.csv':
                self.open_csv(filepath)
            elif ext == '.docx':
                self.open_docx(filepath)
            elif ext == '.pdf':
                self.open_pdf(filepath)
            else:
                # Open as text
                try:
                    with open(filepath, "r", encoding="utf-8") as file:
                        content = file.read()
                        self.text_area.delete("1.0", END)
                        self.text_area.insert("1.0", content)
                        self.notebook.select(0)
                        self.current_file_type = "text"
                        self.file_type_label.config(text="📝 Text Document")
                except Exception as e:
                    messagebox.showerror("Error", f"Could not open file:\n{e}")

            self.is_saved = True
            self.root.title(f"EduText - {os.path.basename(filepath)}")
            self.update_status()

    def save_file(self):
        if self.filename:
            try:
                if self.current_file_type == "csv":
                    with open(self.filename, 'w', newline='', encoding='utf-8') as file:
                        writer = csv.writer(file)
                        writer.writerows(self.csv_data)
                elif self.current_file_type == "docx":
                    self.save_as_docx(self.filename)
                else:
                    content = self.text_area.get("1.0", "end-1c")
                    with open(self.filename, "w", encoding="utf-8") as file:
                        file.write(content)

                self.is_saved = True
                self.root.title(f"EduText - {os.path.basename(self.filename)}")
                messagebox.showinfo("Success", "✅ File saved successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Could not save file:\n{e}")
        else:
            self.save_as_file()

    def save_as_file(self):
        filepath = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[
                ("Text Files", "*.txt"),
                ("CSV Files", "*.csv"),
                ("Word Documents", "*.docx"),
                ("All Files", "*.*")
            ]
        )

        if filepath:
            self.filename = filepath
            self.save_file()

    def export_to_docx(self):
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )

        if filepath:
            self.save_as_docx(filepath)
            messagebox.showinfo("Success", "✅ Exported to DOCX successfully!")

    def save_as_docx(self, filepath):
        """Save text content as DOCX"""
        try:
            doc = Document()
            content = self.text_area.get("1.0", "end-1c")

            for line in content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)

            doc.save(filepath)
        except Exception as e:
            messagebox.showerror("Error", f"Could not export to DOCX:\n{e}")

    def export_to_pdf(self):
        messagebox.showinfo(
            "Export to PDF",
            "📑 To export as PDF:\n\n"
            "1️⃣ Export to DOCX first\n"
            "2️⃣ Open in Word/LibreOffice\n"
            "3️⃣ Save as PDF\n\n"
            "💡 Or use 'Print to PDF' option"
        )

    def on_text_change(self, event=None):
        if self.text_area.edit_modified():
            self.is_saved = False
            self.update_title()
            self.text_area.edit_modified(False)

    def update_title(self):
        title = self.root.title().rstrip(" •")
        if not self.is_saved:
            # add the dot marker
            if not title.endswith(" •"):
                self.root.title(title + " •")

    def update_status(self, event=None):
        content = self.text_area.get("1.0", "end-1c")
        lines = content.count("\n") + 1
        chars = len(content)
        words = len(content.split()) if content.strip() else 0

        self.status_bar.config(
            text=f"📄 Lines: {lines}  |  Characters: {chars}  |  Words: {words}  |  ✓ Ready"
        )

    def on_closing(self):
        if not self.is_saved:
            response = messagebox.askyesnocancel("Save Changes", "💾 Save changes before closing?")
            if response is None:
                return
            elif response:
                self.save_file()
        self.root.destroy()


def textEditor_launch():
    app_window.attributes("-alpha", 1.0)
    app = ModernEditorApp(app_window)
    app_window.mainloop()

# =============================================================================
# CUSTOM DIALOG BOXES
# =============================================================================
class StartingWindow(Toplevel):
    """Starting Screen"""
    def __init__(self):
        super().__init__()
        self.result = None

        self.geometry("960x540")
        self.title("Mail Merge")
        self.resizable(True, True)
        self.attributes("-topmost", True)
        self.overrideredirect(True) # Removes Title Bar

        self.place_window_center()

        self.photo = PhotoImage(file="logo.png")
        image_label = Label(self, image=self.photo)
        image_label.place(x=250, y=45)

        text = ttk.Label(self, text="EduText: A Student Software", font=("Segoe UI", 20, "bold"))
        text.pack(side = "top", pady = 20)
        text2 = ttk.Label(self, text="By: Akshaj Goel", font=("Segoe UI", 14))
        text2.pack(side="bottom", pady=20)

        buttons = ["📧Mail Merge", "📝Text Editor", "❌Exit"]

        bottom_frame = ttk.Frame(self)
        bottom_frame.place(relx=0, rely=1, anchor="sw", relwidth=1)

        for text in buttons:
            btn1 = ttk.Button(bottom_frame, text=text, command=lambda v=text: self.finish(v))

            btn1.pack(side="left", expand=True, fill="x", padx=10, pady=10)

        self.bind("<Button-1>", self.start_move)
        self.bind("<B1-Motion>", self.do_move)

        self.mainloop()

    def start_move(self, event):
        self._x = event.x
        self._y = event.y

    def do_move(self, event):
        dx = event.x - self._x
        dy = event.y - self._y
        x = self.winfo_x() + dx
        y = self.winfo_y() + dy
        self.geometry(f"+{x}+{y}")

    def finish(self, value):
        self.result = value
        self.destroy()
        self.quit()

class CustomBox(Toplevel):
    def __init__(self, buttons, message, image_y):
        super().__init__()
        self.result = None

        self.geometry("480x350")
        self.title("Mail Merge")
        self.resizable(True, True)
        self.attributes("-topmost", True)
        self.overrideredirect(True)

        self.photo = PhotoImage(file="label.png")
        image_label = Label(self, image=self.photo)
        image_label.place(x=128, y= image_y)

        bottom_frame = ttk.Frame(self)
        bottom_frame.place(relx=0, rely=1, anchor="sw", relwidth=1)

        for text in buttons:
            btn1 = ttk.Button(bottom_frame, text=text, command=lambda v=text: self.finish(v))

            btn1.pack(side="left", expand=True, fill="x", padx=5, pady=5)

        text = ttk.Label(self, text=message, font=("Segoe UI", 10))
        text.pack(side="top", pady = 15)

        self.place_window_center()

        self.bind("<Button-1>", self.start_move)
        self.bind("<B1-Motion>", self.do_move)

        self.mainloop()

    def start_move(self, event):
        self._x = event.x
        self._y = event.y

    def do_move(self, event):
        dx = event.x - self._x
        dy = event.y - self._y
        x = self.winfo_x() + dx
        y = self.winfo_y() + dy
        self.geometry(f"+{x}+{y}")

    def finish(self, value):
        self.result = value
        self.destroy()
        self.quit()

class IntBox(Toplevel):
    def __init__(self, buttons, message):
        super().__init__()

        self.result = None
        self.choice = ""

        self.spinbox = ttk.Spinbox(
            self,
            from_=2,
            to=1000,
            width=10
        )
        self.spinbox.pack(side="bottom", pady=60)


        self.geometry("580x400")
        self.title("Mail Merge")
        self.resizable(True, True)
        self.attributes("-topmost", True)
        self.overrideredirect(True)

        self.photo = PhotoImage(file="label.png")
        image_label = Label(self, image=self.photo)
        image_label.place(x=168, y=40)

        bottom_frame = ttk.Frame(self)
        bottom_frame.place(relx=0, rely=1, anchor="sw", relwidth=1)

        for text in buttons:
            btn1 = ttk.Button(bottom_frame, text=text, command=lambda v=text: self.finish(v))

            btn1.pack(side="left", expand=True, fill="x", padx=5, pady=5)

        text = ttk.Label(self, text=message, font=("Segoe UI", 10))
        text.place(relx=0.5, y=20, anchor='n')

        self.place_window_center()

        self.bind("<Button-1>", self.start_move)
        self.bind("<B1-Motion>", self.do_move)

        self.mainloop()

    def start_move(self, event):
        self._x = event.x
        self._y = event.y

    def do_move(self, event):
        dx = event.x - self._x
        dy = event.y - self._y
        x = self.winfo_x() + dx
        y = self.winfo_y() + dy
        self.geometry(f"+{x}+{y}")

    def finish(self, value):
        self.result = int(self.spinbox.get())
        self.choice = value
        self.destroy()
        self.quit()


class StringBox(Toplevel):
    def __init__(self, buttons, message):
        super().__init__()

        self.result = None
        self.choice = ""

        self.entry = ttk.Entry(self)
        self.entry.pack(side="bottom", pady=60)

        self.geometry("580x400")
        self.title("Mail Merge")
        self.resizable(True, True)
        self.attributes("-topmost", True)
        self.overrideredirect(True)

        self.photo = PhotoImage(file="label.png")
        image_label = Label(self, image=self.photo)
        image_label.place(x=168, y=40)

        bottom_frame = ttk.Frame(self)
        bottom_frame.place(relx=0, rely=1, anchor="sw", relwidth=1)

        for text in buttons:
            btn1 = ttk.Button(bottom_frame, text=text, command=lambda v=text: self.finish(v))

            btn1.pack(side="left", expand=True, fill="x", padx=5, pady=5)

        text = ttk.Label(self, text=message, font=("Segoe UI", 10))
        text.place(relx=0.5, y=20, anchor='n')

        self.place_window_center()

        self.bind("<Button-1>", self.start_move)
        self.bind("<B1-Motion>", self.do_move)

        self.mainloop()

    def start_move(self, event):
        self._x = event.x
        self._y = event.y

    def do_move(self, event):
        dx = event.x - self._x
        dy = event.y - self._y
        x = self.winfo_x() + dx
        y = self.winfo_y() + dy
        self.geometry(f"+{x}+{y}")

    def finish(self, value):
        self.result = self.entry.get()
        self.choice = value
        self.destroy()
        self.quit()

class TextBox(Toplevel):
    def __init__(self, buttons, message):
        super().__init__()

        self.result = None
        self.choice = ""

        self.text_box = ttk.Text(self, width=300, height=100)
        self.text_box.pack(padx=50, pady=50)

        self.geometry("960x540")
        self.title("Mail Merge")
        self.resizable(True, True)
        self.attributes("-topmost", True)
        self.overrideredirect(True)

        bottom_frame = ttk.Frame(self)
        bottom_frame.place(relx=0, rely=1, anchor="sw", relwidth=1)

        for text in buttons:
            btn1 = ttk.Button(bottom_frame, text=text, command=lambda v=text: self.finish(v))

            btn1.pack(side="left", expand=True, fill="x", padx=5, pady=5)

        text = ttk.Label(self, text=message, font=("Segoe UI", 10))
        text.place(relx=0.5, y=20, anchor='n')

        self.place_window_center()

        self.bind("<Button-1>", self.start_move)
        self.bind("<B1-Motion>", self.do_move)

        self.mainloop()

    def start_move(self, event):
        self._x = event.x
        self._y = event.y

    def do_move(self, event):
        dx = event.x - self._x
        dy = event.y - self._y
        x = self.winfo_x() + dx
        y = self.winfo_y() + dy
        self.geometry(f"+{x}+{y}")

    def finish(self, value):
        self.result = self.get_text()
        self.choice = value
        self.destroy()
        self.quit()

    def get_text(self):
        content = self.text_box.get("1.0", "end-1c")  # '1.0' = start, 'end-1c' = remove final newline
        return content




# =============================================================================
# LETTER CLASS
# =============================================================================
class Letter:
    def __init__(self):
        self.letter_body = ""
        self.output_directory = ""
        self.letter_file_location = ""
        self.is_placeholder_not_present = False

    def show_placeholder_instructions(self):
        """Displays instructions about using the [name] placeholder."""
        Messagebox.show_info(title="Mail Merge",
                             message="Make sure that you've entered the [name] placeholder in your letter")

    def generate_personalized_letters(self):
        """
        Creates personalized letter files for each recipient.
        Replaces [name] placeholder with actual names and saves files.
        """

        CustomBox(buttons=["Continue", "Exit"],
                  message="Now please select the folders in which you want to save your mails.",
                  image_y=40
                  )

        self.output_directory = filedialog.askdirectory(
            title="Where do you want to save the mails?"
        )

        # Create individual letter file for each recipient
        for recipient in name_manager.recipient_names:
            try:
                output_file = open(f"{self.output_directory}/{recipient}'s Mail.txt", "r")
                output_file.close()

                with open(f"{self.output_directory}/{recipient}'s Mail(1).txt", "r") as output_file:
                    personalized_content = self.letter_body.replace("[name]", recipient)
                    output_file.write(personalized_content)

            except FileNotFoundError:
                with open(f"{self.output_directory}/{recipient}'s Mail.txt", "r") as output_file:
                    personalized_content = self.letter_body.replace("[name]", recipient)
                    output_file.write(personalized_content)

        CustomBox(["Exit"], "Your Mail Merge was successful, congratulations!", 40)

        # Open mail folder
        if platform.system() == "Windows":
            os.startfile(self.output_directory)
        elif platform.system() == "Darwin":  # macOS
            subprocess.call(["open", self.output_directory])
        else:
            subprocess.call(["xdg-open", self.output_directory])

    def process_letter_content(self):
        """
        Handles letter content input based on user's choice (Browse or Type).
        Validates placeholder presence and triggers letter generation.
        """

        letter_choice_dialog = CustomBox(["Browse", "Type Letter Content"],
                                         "How do you want to enter the letter content?\nAs a file or type it?",
                                         50)
        file_input_choice = letter_choice_dialog.result

        if file_input_choice is None:
            show_exit_confirmation(self.process_letter_content)
            return

        if file_input_choice == "Browse":
            self.show_placeholder_instructions()

            # Get letter file from user
            self.letter_file_location = filedialog.askopenfilename(
                title="Select your mail",
                filetypes=(("Text files", "*.txt"),)
            )

            if not self.letter_file_location:
                show_exit_confirmation(self.process_letter_content)
                return

            with open(self.letter_file_location) as letter_file:
                self.letter_body = letter_file.read()

                # Validate [name] placeholder exists
                name_place = self.letter_body.find("[name]")
                if name_place < 0:
                    CustomBox(["Try Again"], "Warning: Placeholder '[name]' not present in file.", 40)
                    self.is_placeholder_not_present = True
                else:
                    self.is_placeholder_not_present = False

                if self.is_placeholder_not_present:
                    # Re-prompt for file input method
                    self.process_letter_content()
                    return

                self.generate_personalized_letters()

        elif file_input_choice == "Type Letter Content":
            self.show_placeholder_instructions()

            # Get letter content via text input
            letter_body_dialog = TextBox(["Continue", "Exit"], "Enter the letter Content")
            self.letter_body = letter_body_dialog.result
            continue_choice = letter_body_dialog.choice


            if continue_choice == "Exit":
                show_exit_confirmation(self.process_letter_content)
                return

            elif continue_choice == "":
                Messagebox.show_warning(title="Mail Merge", message="The letter is empty, please try again.")
                self.process_letter_content()
                return

            # Validate [name] placeholder exists
            name_place = self.letter_body.find("[name]")
            if name_place < 0:
                Messagebox.show_warning(
                    title="Mail Merge",
                    message="Warning: Placeholder '[name]' is not in the letter body"
                )
                self.is_placeholder_not_present = True
            else:
                self.is_placeholder_not_present = False

            if self.is_placeholder_not_present:
                # Re-prompt for file input method
                self.process_letter_content()
                return

            self.generate_personalized_letters()


# =============================================================================
# NAME MANAGER CLASS
# =============================================================================
class NameManager:
    def __init__(self):
        self.recipient_names = []
        self.name_input_method = ""
        self.names_file_location = ""
        self.name_input_method_prompt = "Do you want to enter names manually, or through a text file?"
        self.recipient_count_prompt = "How many people do you want send a mail to?"
        self.name_prompt = "Please enter the name:"
        self.total_recipients = 0

    def name_collection(self):
        """Ask user for name input method"""
        while True:
            name_input_method_dialog = CustomBox(buttons=["Manual Insert", "Text File", "Exit"],
                                                 message=self.name_input_method_prompt,
                                                 image_y=40)
            self.name_input_method = name_input_method_dialog.result

            if self.name_input_method is None:
                exit_confirmation()

            if self.name_input_method == "Manual Insert":
                self.ask_recipient_count()
                self.collect_names_manually()

            elif self.name_input_method == "Text File":
                self.names_text_file_processing()

            else:
                exit_confirmation()

            self.confirm_names_proceed()
            break


    def ask_recipient_count(self):
        """Prompts user to enter the number of recipients."""
        while True:
            total_recipients_dialog_box = IntBox(["Continue", "Exit"], self.recipient_count_prompt)
            self.total_recipients = total_recipients_dialog_box.result

            # Handle cancellation of recipient count dialog
            if total_recipients_dialog_box.choice == "Exit":
                self.total_recipients = 0
                exit_confirmation()

            if self.total_recipients >= 2:
                break

    def collect_names_manually(self):
        """
        Collects names from user through dialog boxes.
        Validates that each name contains only alphabetic characters, spaces, and hyphens.
        """

        for _ in range(self.total_recipients - len(self.recipient_names)):
            while True:
                entered_name_dialog = StringBox(["Continue", "Exit"], self.name_prompt)
                entered_name = entered_name_dialog.result
                name_dialog_exit_choice = entered_name_dialog.choice

                # Handle dialog cancellation
                if name_dialog_exit_choice == "Exit":
                    show_exit_confirmation(self.collect_names_manually)
                    return

                # Validate name contains only letters, spaces, and hyphens
                if entered_name.replace(" ", "").replace("-", "").isalpha():
                    self.name_prompt = "Please enter the name:"
                    self.recipient_names.append(entered_name.title())
                    break
                else:
                    self.name_prompt = "Please enter a valid name:"

    def ask_name_file(self):
        while True:
            self.names_file_location = ""
            self.names_file_location = filedialog.askopenfilename(
                title="Select the names text file",
                filetypes=(("Text files", "*.txt"),)
            )

            if not self.names_file_location:
                exit_confirmation()
            else:
                break

    def names_text_file_processing(self):
        """Get names file from user"""
        while True:
            self.ask_name_file()

            try:
                with open(self.names_file_location) as names_file:
                    names_content = names_file.read()

            except Exception as e:
                Messagebox.show_warning(title = "Error occurred",
                                        message = f"Error: {e}")
                self.name_collection()
                return

            # Validate comma separation
            if ", " not in names_content:
                Messagebox.show_warning(
                    title="Mail Merge",
                    message="Please enter the names separated by comma."
                )

            elif ", " in names_content:
                break

        self.recipient_names = names_content.split(", ")

    def confirm_names_proceed(self):
        """
        Shows collected names to user for confirmation.
        Allows re-entry or continuation to file selection.
        Returns user's file input choice (Browse or Type).
        """
        global confirmation_choice

        names_confirmation_dialog = CustomBox(
            message=f"Names entered:\n{', '.join(self.recipient_names)}",
            buttons=["Continue", "Re-enter Names", "Exit"],
            image_y=60
        )

        confirmation_choice = names_confirmation_dialog.result

        if confirmation_choice == "Continue":
            return confirmation_choice

        elif confirmation_choice == "Re-enter Names":
            self.recipient_names = []
            self.name_collection()
            return
        else:
            exit_confirmation()

# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def center_window(window):
    """Centers the given window on the screen."""
    window.update_idletasks()
    width = window.winfo_width()
    height = window.winfo_height()
    x = (window.winfo_screenwidth() // 2) - (width // 2)
    y = (window.winfo_screenheight() // 2) - (height // 2)
    window.geometry(f'{width}x{height}+{x}+{y}')


def show_exit_confirmation(callback_function):
    """
    Shows exit confirmation dialog and handles user's choice.
    If user chooses 'No', calls the provided callback function.
    """
    exit_dialog = CustomBox(["Yes", "No"], "Do you want to exit the app?", 40)

    user_choice = exit_dialog.result

    if user_choice == "Yes":
        exit()
    else:
        callback_function()

def exit_confirmation():
    """
    Shows exit confirmation dialog and handles user's choice.
    If user chooses 'No', calls the provided callback function.
    """
    exit_dialog = CustomBox(["Yes", "No"], "Do you want to exit the app?", 40)

    user_choice = exit_dialog.result

    if user_choice == "Yes":
        exit()
    else:
        pass


# =============================================================================
# INITIALIZATION & MAIN FLOW
# =============================================================================
# Loading Screen
app_window = ttk.Window(themename="darkly")
app_window.attributes('-alpha', 0)
app_window.iconbitmap("logo.ico")
center_window(app_window)

# Initialize variables
confirmation_choice = ""

# Ask user for name input method
input_method_choice = ""

# Objects
name_manager = NameManager()
letter = Letter()

while True:
    start_screen = StartingWindow()
    software_choice = start_screen.result
    if software_choice == "📧Mail Merge":
        break
    elif software_choice == "📝Text Editor":
        break
    elif software_choice == "❌Exit":
        exit_confirmation()

if software_choice == "📧Mail Merge":
    name_manager.name_collection()
    letter.process_letter_content()

elif software_choice == "📝Text Editor":
    textEditor_launch()