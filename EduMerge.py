# =============================================================================
# IMPORTS
# =============================================================================
# GUI and Dialog imports
from tkinter import filedialog, messagebox, colorchooser
from tkinter import Text, END, LEFT, RIGHT, BOTTOM, TOP, X, Y, BOTH, W, E
import ttkbootstrap as ttk
from ttkbootstrap import Toplevel, Label, PhotoImage
from ttkbootstrap.dialogs import Messagebox
from ttkbootstrap.widgets.tableview import Tableview        # fixed deprecated import
from ttkbootstrap.constants import *
from pathlib import Path

# System imports
import os, platform, subprocess

# Image processing
from PIL import Image, ImageTk

# File format handling
import csv
from docx import Document
from docx.shared import Pt, RGBColor
import pdf2image


# =============================================================================
# Text Editor
# =============================================================================
class ModernEditorApp:
    """
    Main text editor application with modern UI.
    Supports multiple file formats: TXT, CSV, DOCX, PDF
    Features: Rich text formatting, image insertion, export capabilities
    """

    def __init__(self, root):
        """
        Initialize the text editor application.

        Args:
            root: The main tkinter window
        """
        self.root = root
        self.root.title("EduText: Text Editor")
        self.root.geometry("1200x800")

        # Set application icon (fails silently if file not found)
        try:
            self.root.iconbitmap("logo.ico")
        except:
            pass

        # Application state variables
        self.filename = None  # Current file path
        self.is_saved = True  # Track if document has unsaved changes
        self.images = []  # Store PhotoImage objects to prevent garbage collection
        self.current_file_type = "text"  # Current document type (text/csv/docx/pdf)

        # CSV-specific data storage
        self.csv_data = []  # Raw CSV data rows
        self.csv_headers = []  # CSV column headers

        # Apply custom styling
        self.apply_custom_styles()

        # Build the UI components
        self.create_modern_sidebar()
        self.create_main_area()
        self.create_status_bar()

        # Register keyboard shortcuts
        self.root.bind("<Control-s>", lambda e: self.save_file())
        self.root.bind("<Control-o>", lambda e: self.open_file())
        self.root.bind("<Control-n>", lambda e: self.new_file())
        self.root.bind("<Control-b>", lambda e: self.apply_bold())
        self.root.bind("<Control-i>", lambda e: self.apply_italic())
        self.root.bind("<Control-u>", lambda e: self.apply_underline())

        # Handle window close event
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def apply_custom_styles(self):
        """
        Apply custom ttk styles for buttons and UI elements.
        Uses safe style configurations without unsupported builder methods.
        """
        style = ttk.Style()

        # Configure base button style with padding and flat appearance
        style.configure('TButton', padding=(12, 8), borderwidth=0, relief="flat")

        # Create smaller button variant
        style.configure('Small.TButton', padding=(8, 6), borderwidth=0)

    def create_modern_sidebar(self):
        """
        Create the left sidebar containing app logo, navigation, and action buttons.
        Organized into sections: File Operations, Insert Content, Export Options.
        """
        # Main sidebar container with dark theme
        self.sidebar = ttk.Frame(self.root, bootstyle="dark", width=280)
        self.sidebar.pack(side=LEFT, fill=Y, padx=0, pady=0)
        self.sidebar.pack_propagate(False)  # Prevent sidebar from shrinking

        # === LOGO/TITLE SECTION ===
        title_frame = ttk.Frame(self.sidebar, bootstyle="dark")
        title_frame.pack(fill=X, padx=25, pady=35)

        # App icon (emoji)
        ttk.Label(
            title_frame,
            text="📝",
            font=("Segoe UI Emoji", 32),
            bootstyle="light"
        ).pack()

        # App title
        ttk.Label(
            title_frame,
            text="EduText",
            font=("Segoe UI", 22, "bold"),
            bootstyle="light",
            foreground="#FFFFFF"
        ).pack(pady=(8, 2))

        # Subtitle
        ttk.Label(
            title_frame,
            text="A Student Software",
            font=("Segoe UI", 9),
            bootstyle="light",
            foreground="#A0A0A0"
        ).pack()

        # Visual separator
        ttk.Separator(self.sidebar, bootstyle="secondary").pack(fill=X, padx=25, pady=15)

        # === FILE OPERATIONS SECTION ===
        self.create_section_header("FILE OPERATIONS")
        self.create_sidebar_button("📄  New Document", self.new_file, "success")
        self.create_sidebar_button("📂  Open File", self.open_file, "info")
        self.create_sidebar_button("💾  Save", self.save_file, "primary")
        self.create_sidebar_button("💾  Save As", self.save_as_file, "primary")

        # === INSERT CONTENT SECTION ===
        self.create_section_header("INSERT CONTENT")
        self.create_sidebar_button("🖼️  Insert Image", self.insert_image, "warning")
        self.create_sidebar_button("📊  CSV Table", self.create_csv_table, "info")

        # === EXPORT OPTIONS SECTION ===
        self.create_section_header("EXPORT OPTIONS")
        self.create_sidebar_button("📄  Export DOCX", self.export_to_docx, "danger")
        self.create_sidebar_button("📕  Export PDF", self.export_to_pdf, "danger")

        # === FOOTER ===
        footer = ttk.Frame(self.sidebar, bootstyle="dark")
        footer.pack(side=BOTTOM, fill=X, padx=25, pady=20)

        # Version information
        ttk.Label(
            footer,
            text="v0.2 • Modern Edition",
            font=("Segoe UI", 8),
            bootstyle="light",
            foreground="#707070"
        ).pack()

    def create_section_header(self, text):
        """
        Create a styled section header for the sidebar.

        Args:
            text: Header text to display
        """
        header_frame = ttk.Frame(self.sidebar, bootstyle="dark")
        header_frame.pack(fill=X, padx=25, pady=(20, 12))

        ttk.Label(
            header_frame,
            text=text,
            font=("Segoe UI", 8, "bold"),
            bootstyle="light",
            foreground="#808080"
        ).pack(anchor=W)

    def create_sidebar_button(self, text, command, style):
        """
        Create a styled button in the sidebar.

        Args:
            text: Button label text
            command: Function to call when button is clicked
            style: Bootstrap style variant (e.g., "success", "info", "primary")
        """
        btn = ttk.Button(
            self.sidebar,
            text=text,
            command=command,
            bootstyle=style,
            width=28,
            cursor="hand2"  # Hand cursor on hover
        )
        btn.pack(padx=25, pady=4, fill=X)

    def create_main_area(self):
        """
        Create the main content area containing toolbar, text editor, and tabbed views.
        Includes tabs for: Text Editor, CSV Table, PDF Viewer.
        """
        # Main container for content area
        main_container = ttk.Frame(self.root, bootstyle="light")
        main_container.pack(side=LEFT, fill=BOTH, expand=True)

        # Create toolbar with formatting controls
        self.create_modern_toolbar(main_container)

        # === CONTENT WRAPPER ===
        content_wrapper = ttk.Frame(main_container, bootstyle="light")
        content_wrapper.pack(fill=BOTH, expand=True, padx=25, pady=(15, 20))

        # Card-style container
        card_frame = ttk.Frame(content_wrapper, bootstyle="light", relief="flat")
        card_frame.pack(fill=BOTH, expand=True)

        # === TABBED NOTEBOOK ===
        self.notebook = ttk.Notebook(card_frame, bootstyle="dark")
        self.notebook.pack(fill=BOTH, expand=True)

        # === TEXT EDITOR TAB ===
        self.text_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.text_frame, text="✏️  Text Editor")

        # Main text widget with modern styling
        self.text_area = Text(
            self.text_frame,
            wrap="word",  # Word wrapping
            font=("Segoe UI", 12),
            bg="#FFFFFF",
            fg="#1A1A1A",
            insertbackground="#3498DB",  # Cursor color
            relief="flat",
            padx=30,
            pady=30,
            selectbackground="#3498DB",  # Selection highlight color
            selectforeground="white",
            spacing1=3,  # Space above lines
            spacing3=3  # Space below lines
        )

        # Scrollbar for text area
        text_scrollbar = ttk.Scrollbar(self.text_frame, command=self.text_area.yview, bootstyle="round")
        self.text_area.configure(yscrollcommand=text_scrollbar.set)

        text_scrollbar.pack(side=RIGHT, fill=Y)
        self.text_area.pack(fill=BOTH, expand=True)

        # === TEXT FORMATTING TAGS ===
        # Configure custom text styles
        self.text_area.tag_configure("bold", font=("Segoe UI", 12, "bold"))
        self.text_area.tag_configure("italic", font=("Segoe UI", 12, "italic"))
        self.text_area.tag_configure("underline", font=("Segoe UI", 12, "underline"))
        self.text_area.tag_configure("heading1", font=("Segoe UI", 32, "bold"), foreground="#E74C3C", spacing3=15)
        self.text_area.tag_configure("heading2", font=("Segoe UI", 24, "bold"), foreground="#3498DB", spacing3=12)
        self.text_area.tag_configure("heading3", font=("Segoe UI", 18, "bold"), foreground="#16A085", spacing3=10)

        # Bind events for tracking changes and updating status
        self.text_area.bind("<<Modified>>", self.on_text_change)
        self.text_area.bind("<KeyRelease>", self.update_status)

        # === CSV TABLE TAB ===
        self.csv_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.csv_frame, text="📊  CSV Table")

        # === PDF VIEWER TAB ===
        self.pdf_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.pdf_frame, text="📕  PDF Viewer")

        # Read-only text widget for PDF content
        self.pdf_text = Text(
            self.pdf_frame,
            wrap="word",
            font=("Segoe UI", 11),
            bg="#FAFAFA",
            fg="#1A1A1A",
            relief="flat",
            padx=30,
            pady=30,
            state="disabled"  # Read-only
        )

        # PDF scrollbar
        pdf_scrollbar = ttk.Scrollbar(self.pdf_frame, command=self.pdf_text.yview, bootstyle="round")
        self.pdf_text.configure(yscrollcommand=pdf_scrollbar.set)
        pdf_scrollbar.pack(side=RIGHT, fill=Y)
        self.pdf_text.pack(fill=BOTH, expand=True)

    def create_modern_toolbar(self, parent):
        """
        Create toolbar with font controls, text styling, headings, and color options.

        Args:
            parent: Parent widget to attach toolbar to
        """
        toolbar_container = ttk.Frame(parent, bootstyle="light")
        toolbar_container.pack(fill=X, padx=25, pady=(25, 5))

        toolbar = ttk.Frame(toolbar_container, bootstyle="light", relief="flat")
        toolbar.pack(fill=X, pady=5)

        # === FONT CONTROLS CARD ===
        font_card = ttk.Labelframe(
            toolbar,
            text="  Font  ",
            bootstyle="primary",
            padding=15
        )
        font_card.pack(side=LEFT, padx=(0, 10))

        # Font family dropdown
        self.font_family_var = ttk.StringVar(value="Segoe UI")
        fonts = ["Segoe UI", "Arial", "Times New Roman", "Calibri", "Georgia", "Courier New", "Consolas"]
        font_combo = ttk.Combobox(
            font_card,
            textvariable=self.font_family_var,
            values=fonts,
            width=14,
            state="readonly",
            bootstyle="primary"
        )
        font_combo.pack(side=LEFT, padx=(0, 8))
        font_combo.bind("<<ComboboxSelected>>", lambda e: self.change_font())

        # Font size spinner
        self.font_size_var = ttk.IntVar(value=12)
        size_spinbox = ttk.Spinbox(
            font_card,
            from_=8,
            to=72,
            textvariable=self.font_size_var,
            width=6,
            command=self.change_font,
            bootstyle="primary"
        )
        size_spinbox.pack(side=LEFT, padx=2)

        # === TEXT STYLE CARD ===
        style_card = ttk.Labelframe(
            toolbar,
            text="  Style  ",
            bootstyle="success",
            padding=15
        )
        style_card.pack(side=LEFT, padx=5)

        # Bold button
        btn_bold = ttk.Button(
            style_card,
            text="𝐁",
            bootstyle="success-outline",
            command=self.apply_bold,
            width=4,
            cursor="hand2"
        )
        btn_bold.pack(side=LEFT, padx=2)

        # Italic button
        btn_italic = ttk.Button(
            style_card,
            text="𝐼",
            bootstyle="success-outline",
            command=self.apply_italic,
            width=4,
            cursor="hand2"
        )
        btn_italic.pack(side=LEFT, padx=2)

        # Underline button
        btn_underline = ttk.Button(
            style_card,
            text="U̲",
            bootstyle="success-outline",
            command=self.apply_underline,
            width=4,
            cursor="hand2"
        )
        btn_underline.pack(side=LEFT, padx=2)

        # === HEADING CARD ===
        heading_card = ttk.Labelframe(
            toolbar,
            text="  Headings  ",
            bootstyle="info",
            padding=15
        )
        heading_card.pack(side=LEFT, padx=5)

        # H1, H2, H3 buttons with different color styles
        ttk.Button(
            heading_card,
            text="H1",
            bootstyle="danger",
            command=lambda: self.apply_heading("heading1"),
            width=4,
            cursor="hand2"
        ).pack(side=LEFT, padx=2)

        ttk.Button(
            heading_card,
            text="H2",
            bootstyle="primary",
            command=lambda: self.apply_heading("heading2"),
            width=4,
            cursor="hand2"
        ).pack(side=LEFT, padx=2)

        ttk.Button(
            heading_card,
            text="H3",
            bootstyle="success",
            command=lambda: self.apply_heading("heading3"),
            width=4,
            cursor="hand2"
        ).pack(side=LEFT, padx=2)

        # === COLOR CARD ===
        color_card = ttk.Labelframe(
            toolbar,
            text="  Colors  ",
            bootstyle="warning",
            padding=15
        )
        color_card.pack(side=LEFT, padx=5)

        # Text color button
        ttk.Button(
            color_card,
            text="🎨 Text",
            bootstyle="warning",
            command=self.change_text_color,
            width=9,
            cursor="hand2"
        ).pack(side=LEFT, padx=2)

        # Background color button
        ttk.Button(
            color_card,
            text="🎨 BG",
            bootstyle="warning-outline",
            command=self.change_bg_color,
            width=8,
            cursor="hand2"
        ).pack(side=LEFT, padx=2)

    def create_status_bar(self):
        """
        Create bottom status bar showing document statistics and file type.
        Displays: line count, character count, word count, and current file type.
        """
        status_frame = ttk.Frame(self.root, bootstyle="dark")
        status_frame.pack(fill=X, side=BOTTOM)

        # Left side - document statistics
        self.status_bar = ttk.Label(
            status_frame,
            text="📄 Lines: 1  |  Characters: 0  |  Words: 0  |  Ready",
            anchor=W,
            font=("Segoe UI", 9),
            padding=(15, 8)
        )
        self.status_bar.pack(side=LEFT, fill=X, expand=True)

        # Right side - file type indicator
        self.file_type_label = ttk.Label(
            status_frame,
            text="📝 Text Document",
            anchor=E,
            font=("Segoe UI", 9, "bold"),
            padding=(15, 8)
        )
        self.file_type_label.pack(side=RIGHT)

    # =========================================================================
    # TEXT FORMATTING METHODS
    # =========================================================================

    def apply_bold(self):
        """
        Toggle bold formatting on selected text.
        If text is already bold, removes bold; otherwise applies bold.
        """
        try:
            current_tags = self.text_area.tag_names("sel.first")
            if "bold" in current_tags:
                self.text_area.tag_remove("bold", "sel.first", "sel.last")
            else:
                self.text_area.tag_add("bold", "sel.first", "sel.last")
        except:
            pass  # No selection

    def apply_italic(self):
        """Toggle italic formatting on selected text."""
        try:
            current_tags = self.text_area.tag_names("sel.first")
            if "italic" in current_tags:
                self.text_area.tag_remove("italic", "sel.first", "sel.last")
            else:
                self.text_area.tag_add("italic", "sel.first", "sel.last")
        except:
            pass  # No selection

    def apply_underline(self):
        """Toggle underline formatting on selected text."""
        try:
            current_tags = self.text_area.tag_names("sel.first")
            if "underline" in current_tags:
                self.text_area.tag_remove("underline", "sel.first", "sel.last")
            else:
                self.text_area.tag_add("underline", "sel.first", "sel.last")
        except:
            pass  # No selection

    def apply_heading(self, heading_tag):
        """
        Apply heading style to selected text.
        Removes any existing heading tags before applying new one.

        Args:
            heading_tag: Tag name ("heading1", "heading2", or "heading3")
        """
        try:
            # Remove all heading tags first
            for tag in ["heading1", "heading2", "heading3"]:
                self.text_area.tag_remove(tag, "sel.first", "sel.last")
            # Apply the selected heading
            self.text_area.tag_add(heading_tag, "sel.first", "sel.last")
        except:
            pass  # No selection

    def change_font(self):
        """
        Update font family and size for all text and formatting tags.
        Called when user changes font dropdown or size spinner.
        """
        family = self.font_family_var.get()
        size = self.font_size_var.get()

        # Update base font
        self.text_area.configure(font=(family, size))

        # Update all formatting tag fonts
        self.text_area.tag_configure("bold", font=(family, size, "bold"))
        self.text_area.tag_configure("italic", font=(family, size, "italic"))
        self.text_area.tag_configure("underline", font=(family, size, "underline"))
        self.text_area.tag_configure("heading1", font=(family, size + 20, "bold"))
        self.text_area.tag_configure("heading2", font=(family, size + 12, "bold"))
        self.text_area.tag_configure("heading3", font=(family, size + 6, "bold"))

    def change_text_color(self):
        """
        Open color picker and apply selected color to selected text.
        Creates dynamic tag for the chosen color.
        """
        try:
            color = colorchooser.askcolor(title="Choose Text Color")
            if color[1]:  # color[1] is hex value
                tag_name = f"color_{color[1]}"
                self.text_area.tag_configure(tag_name, foreground=color[1])
                self.text_area.tag_add(tag_name, "sel.first", "sel.last")
        except:
            pass  # No selection or dialog cancelled

    def change_bg_color(self):
        """
        Open color picker and apply selected background color to selected text.
        Creates dynamic tag for the chosen background color.
        """
        try:
            color = colorchooser.askcolor(title="Choose Background Color")
            if color[1]:
                tag_name = f"bg_{color[1]}"
                self.text_area.tag_configure(tag_name, background=color[1])
                self.text_area.tag_add(tag_name, "sel.first", "sel.last")
        except:
            pass  # No selection or dialog cancelled

    # =========================================================================
    # IMAGE AND MEDIA METHODS
    # =========================================================================

    def insert_image(self):
        """
        Open file dialog to select and insert an image into the document.
        Automatically resizes large images to fit within 700px width.
        """
        filepath = filedialog.askopenfilename(
            title="Select Image",
            filetypes=[("Image Files", "*.png *.jpg *.jpeg *.gif *.bmp"), ("All Files", "*.*")]
        )

        if filepath:
            try:
                img = Image.open(filepath)
                max_width = 700

                # Resize if image is too wide
                if img.width > max_width:
                    ratio = max_width / img.width
                    new_height = int(img.height * ratio)
                    # Use appropriate resampling method (newer PIL versions use Image.Resampling)
                    try:
                        resample = Image.Resampling.LANCZOS
                    except AttributeError:
                        resample = Image.LANCZOS
                    img = img.resize((max_width, new_height), resample)

                # Convert to PhotoImage and store reference
                photo = ImageTk.PhotoImage(img)
                self.images.append(photo)  # Prevent garbage collection

                # Insert image at cursor position
                self.text_area.image_create("insert", image=photo)
                self.text_area.insert("insert", "\n")

                # Mark as unsaved
                self.is_saved = False
                self.update_title()
            except Exception as e:
                messagebox.showerror("Error", f"Could not insert image:\n{e}")

    # =========================================================================
    # CSV HANDLING METHODS
    # =========================================================================

    def create_csv_table(self):
        """
        Switch to CSV tab view and prepare for CSV operations.
        Updates file type indicator.
        """
        self.notebook.select(1)  # Select CSV tab (index 1)
        self.current_file_type = "csv"
        self.file_type_label.config(text="📊 CSV Table")

    def open_csv(self, filepath):
        """
        Read and display CSV file in table format.
        Uses Tableview widget with pagination and search functionality.

        Args:
            filepath: Path to the CSV file to open
        """
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                csv_reader = csv.reader(file)
                self.csv_data = list(csv_reader)

                if self.csv_data:
                    # First row becomes headers
                    self.csv_headers = self.csv_data[0]

                    # Clear any existing table
                    for widget in self.csv_frame.winfo_children():
                        widget.destroy()

                    # Prepare column definitions
                    coldata = [{"text": header, "stretch": True} for header in self.csv_headers]
                    rowdata = self.csv_data[1:]  # All rows except header

                    # Create table widget
                    table = Tableview(
                        self.csv_frame,
                        coldata=coldata,
                        rowdata=rowdata,
                        paginated=True,  # Enable pagination
                        searchable=True,  # Enable search functionality
                        bootstyle="primary",
                        height=20
                    )
                    table.pack(fill=BOTH, expand=True, padx=20, pady=20)

                    # Switch to CSV tab
                    self.notebook.select(1)
                    self.current_file_type = "csv"
                    self.file_type_label.config(text="📊 CSV Table")
        except Exception as e:
            messagebox.showerror("Error", f"Could not open CSV:\n{e}")

    # =========================================================================
    # DOCUMENT FORMAT METHODS
    # =========================================================================

    def open_docx(self, filepath):
        """
        Open and display Word document in text editor.
        Extracts plain text from DOCX file (formatting is lost).

        Args:
            filepath: Path to the DOCX file
        """
        try:
            doc = Document(filepath)
            self.text_area.delete("1.0", END)

            # Extract all paragraphs as plain text
            for para in doc.paragraphs:
                self.text_area.insert(END, para.text + "\n")

            self.notebook.select(0)  # Switch to text editor tab
            self.current_file_type = "docx"
            self.file_type_label.config(text="📄 Word Document")
        except Exception as e:
            messagebox.showerror("Error", f"Could not open DOCX:\n{e}")

    def open_pdf(self, filepath):
        try:
            from pdf2image import convert_from_path

            # ⚠ Show "wait" window BEFORE freeze
            wait_win = ttk.Toplevel(self)
            wait_win.title("Loading PDF")
            wait_win.geometry("320x120")
            wait_win.resizable(False, False)

            ttk.Label(
                wait_win,
                text="Please wait...\nThe editor may freeze while the PDF loads.",
                font=("Segoe UI", 11),
                justify="center"
            ).pack(pady=15)

            pbar = ttk.Progressbar(wait_win, mode="indeterminate", length=220)
            pbar.pack(pady=5)
            pbar.start(12)

            # Force Tk to draw the window BEFORE freezing
            wait_win.update()

            # ----------------------------
            #  The part that will freeze
            # ----------------------------
            pages = convert_from_path(filepath, dpi=300)

            # Once loaded, close wait window
            wait_win.destroy()

            # ---- Your existing UI code ----
            for widget in self.pdf_frame.winfo_children():
                widget.destroy()

            canvas = ttk.Canvas(self.pdf_frame, bg="#FAFAFA")
            scrollbar = ttk.Scrollbar(
                self.pdf_frame, orient="vertical",
                command=canvas.yview, bootstyle="round"
            )
            scrollable_frame = ttk.Frame(canvas)

            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )

            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)

            scrollbar.pack(side=RIGHT, fill=Y)
            canvas.pack(side=LEFT, fill=BOTH, expand=True)

            self._display_pdf_pages(pages, scrollable_frame, canvas, wait_win)

            self.notebook.select(2)
            self.current_file_type = "pdf"
            self.file_type_label.config(text="📕 PDF Document")

        except ImportError:
            messagebox.showerror(
                "Missing Dependency",
                "PDF viewing requires 'pdf2image'. Install with:\npip install pdf2image"
            )
        except Exception as e:
            messagebox.showerror("Error", f"Could not open PDF:\n{e}")

    def _display_pdf_pages(self, pages, scrollable_frame, canvas, loading_win):
        loading_win.destroy()

        for page_num, page_image in enumerate(pages, 1):
            max_width = 800
            if page_image.width > max_width:
                ratio = max_width / page_image.width
                new_height = int(page_image.height * ratio)
                try:
                    resample = Image.Resampling.LANCZOS
                except AttributeError:
                    resample = Image.LANCZOS
                page_image = page_image.resize((max_width, new_height), resample)

            photo = ImageTk.PhotoImage(page_image)
            self.images.append(photo)

            lbl = ttk.Label(scrollable_frame,
                            text=f"Page {page_num}",
                            font=("Segoe UI", 12, "bold"),
                            bootstyle="info")
            lbl.pack(pady=(20, 5))

            img_label = ttk.Label(scrollable_frame, image=photo)
            img_label.pack(pady=(0, 10))

            if page_num < len(pages):
                ttk.Separator(scrollable_frame, bootstyle="secondary").pack(fill=X, padx=50, pady=10)

        # Mouse wheel scroll
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        canvas.bind_all("<MouseWheel>", _on_mousewheel)

        self.notebook.select(2)

    # =========================================================================
    # FILE OPERATION METHODS
    # =========================================================================

    def new_file(self):
        """
        Create a new document, prompting to save if there are unsaved changes.
        Resets all document state variables.
        """
        # Check for unsaved changes
        if not self.is_saved:
            response = messagebox.askyesnocancel("Save Changes", "Do you want to save changes?")
            if response is None:
                return  # User cancelled
            elif response:
                self.save_file()

        # Reset editor state
        self.text_area.delete("1.0", END)
        self.images.clear()
        self.csv_data.clear()
        self.filename = None
        self.is_saved = True
        self.current_file_type = "text"
        self.file_type_label.config(text="📝 Text Document")
        self.notebook.select(0)
        self.update_status()

    def open_file(self):
        """
        Open a file with auto-detection of format (TXT, CSV, DOCX, PDF).
        Routes to appropriate handler based on file extension.
        """
        filepath = filedialog.askopenfilename(
            defaultextension=".txt",
            filetypes=[
                ("All Supported", "*.txt *.csv *.docx *.pdf"),
                ("Text Files", "*.txt"),
                ("CSV Files", "*.csv"),
                ("Word Documents", "*.docx"),
                ("PDF Files", "*.pdf"),
                ("All Files", "*.*")
            ]
        )

        if filepath:
            self.filename = filepath
            ext = os.path.splitext(filepath)[1].lower()

            # Route to appropriate handler based on extension
            if ext == '.csv':
                self.open_csv(filepath)
            elif ext == '.docx':
                self.open_docx(filepath)
            elif ext == '.pdf':
                self.open_pdf(filepath)
            else:
                # Open as plain text
                try:
                    with open(filepath, "r", encoding="utf-8") as file:
                        content = file.read()
                        self.text_area.delete("1.0", END)
                        self.text_area.insert("1.0", content)
                        self.notebook.select(0)
                        self.current_file_type = "text"
                        self.file_type_label.config(text="📝 Text Document")
                except Exception as e:
                    messagebox.showerror("Error", f"Could not open file:\n{e}")

            # Update window state
            self.is_saved = True
            self.root.title(f"EduText - {os.path.basename(filepath)}")
            self.update_status()

    def save_file(self):
        """
        Save the current document to disk.
        Handles different file types appropriately (CSV, DOCX, TXT).
        If no filename exists, prompts for Save As.
        """
        if self.filename:
            try:
                # Save based on current file type
                if self.current_file_type == "csv":
                    with open(self.filename, 'w', newline='', encoding='utf-8') as file:
                        writer = csv.writer(file)
                        writer.writerows(self.csv_data)
                elif self.current_file_type == "docx":
                    self.save_as_docx(self.filename)
                else:
                    # Save as plain text
                    content = self.text_area.get("1.0", "end-1c")
                    with open(self.filename, "w", encoding="utf-8") as file:
                        file.write(content)

                # Mark as saved
                self.is_saved = True
                self.root.title(f"EduText - {os.path.basename(self.filename)}")
                messagebox.showinfo("Success", "✅ File saved successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Could not save file:\n{e}")
        else:
            # No filename set, prompt for Save As
            self.save_as_file()

    def save_as_file(self):
        """
        Prompt user for new filename and save document.
        Allows user to choose file format (TXT, CSV, DOCX).
        """
        filepath = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[
                ("Text Files", "*.txt"),
                ("CSV Files", "*.csv"),
                ("Word Documents", "*.docx"),
                ("All Files", "*.*")
            ]
        )

        if filepath:
            self.filename = filepath
            self.save_file()

    # =========================================================================
    # EXPORT METHODS
    # =========================================================================

    def export_to_docx(self):
        """
        Export current document content to Word format (DOCX).
        Prompts for save location and creates formatted Word document.
        """
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )

        if filepath:
            self.save_as_docx(filepath)
            messagebox.showinfo("Success", "✅ Exported to DOCX successfully!")

    def save_as_docx(self, filepath):
        """
        Create Word document from text content.
        Saves each non-empty line as a separate paragraph.

        Args:
            filepath: Destination path for DOCX file
        """
        try:
            doc = Document()
            content = self.text_area.get("1.0", "end-1c")

            # Add each line as a paragraph
            for line in content.split("\n"):
                if line.strip():  # Only add non-empty lines
                    doc.add_paragraph(line)

            doc.save(filepath)
        except Exception as e:
            messagebox.showerror("Error", f"Could not export to DOCX:\n{e}")

    def export_to_pdf(self):
        """
        Display instructions for PDF export.
        Note: Direct PDF export requires additional libraries, so this shows workaround.
        """
        messagebox.showinfo(
            "Export to PDF",
            "📕 To export as PDF:\n\n"
            "1️⃣ Export to DOCX first\n"
            "2️⃣ Open in Word/LibreOffice\n"
            "3️⃣ Save as PDF\n\n"
            "💡 Or use 'Print to PDF' option"
        )

    # =========================================================================
    # EVENT HANDLERS
    # =========================================================================

    def on_text_change(self, event=None):
        """
        Handle text modification events.
        Marks document as unsaved and updates window title.

        Args:
            event: Tkinter event object (optional)
        """
        if self.text_area.edit_modified():
            self.is_saved = False
            self.update_title()
            self.text_area.edit_modified(False)  # Reset modified flag

    def update_title(self):
        """
        Update window title to reflect unsaved changes.
        Adds bullet point (•) indicator when document has unsaved changes.
        """
        title = self.root.title().rstrip(" •")
        if not self.is_saved:
            # Add unsaved indicator
            if not title.endswith(" •"):
                self.root.title(title + " •")

    def update_status(self, event=None):
        """
        Update status bar with current document statistics.
        Displays: line count, character count, word count.

        Args:
            event: Tkinter event object (optional)
        """
        content = self.text_area.get("1.0", "end-1c")
        lines = content.count("\n") + 1
        chars = len(content)
        words = len(content.split()) if content.strip() else 0

        self.status_bar.config(
            text=f"📄 Lines: {lines}  |  Characters: {chars}  |  Words: {words}  |  ✓ Ready"
        )

    def on_closing(self):
        """
        Handle window close event.
        Prompts to save if there are unsaved changes before closing.
        """
        if not self.is_saved:
            response = messagebox.askyesnocancel("Save Changes", "💾 Save changes before closing?")
            if response is None:
                return  # User cancelled, don't close
            elif response:
                self.save_file()
        self.root.destroy()


def textEditor_launch():
    """
    Launch the text editor application.
    Makes window visible and starts main event loop.
    """
    app_window.attributes("-alpha", 1.0)  # Make window visible
    app = ModernEditorApp(app_window)
    app_window.mainloop()


# =============================================================================
# CUSTOM DIALOG BOXES
# =============================================================================

class StartingWindow(Toplevel):
    """
    Custom starting screen for application launcher.
    Allows user to choose between Mail Merge and Text Editor.
    Features: Draggable window, custom styling, no title bar.
    """

    def __init__(self):
        """Initialize the starting window with logo and option buttons."""
        super().__init__()
        self.result = None  # Stores user's choice

        # Window configuration
        self.geometry("960x540")
        self.title("Mail Merge")
        self.resizable(True, True)
        self.attributes("-topmost", True)  # Keep on top
        self.overrideredirect(True)  # Remove title bar for custom look

        # Center window on screen
        self.place_window_center()

        # Display logo
        self.photo = PhotoImage(file="logo.png")
        image_label = Label(self, image=self.photo)
        image_label.place(x=250, y=45)

        # Title text
        text = ttk.Label(self, text="EduText: A Student Software", font=("Segoe UI", 20, "bold"))
        text.pack(side="top", pady=20)

        # Author credit
        text2 = ttk.Label(self, text="By: Akshaj Goel", font=("Segoe UI", 14))
        text2.pack(side="bottom", pady=20)

        # Option buttons
        buttons = ["📧Mail Merge", "📝Text Editor", "❌Exit"]

        # Bottom button bar
        bottom_frame = ttk.Frame(self)
        bottom_frame.place(relx=0, rely=1, anchor="sw", relwidth=1)

        # Create buttons dynamically
        for text in buttons:
            btn1 = ttk.Button(bottom_frame, text=text, command=lambda v=text: self.finish(v))
            btn1.pack(side="left", expand=True, fill="x", padx=10, pady=10)

        # Enable window dragging
        self.bind("<Button-1>", self.start_move)
        self.bind("<B1-Motion>", self.do_move)

        self.mainloop()

    def start_move(self, event):
        """Record initial mouse position for dragging."""
        self._x = event.x
        self._y = event.y

    def do_move(self, event):
        """Calculate and apply window movement based on mouse drag."""
        dx = event.x - self._x
        dy = event.y - self._y
        x = self.winfo_x() + dx
        y = self.winfo_y() + dy
        self.geometry(f"+{x}+{y}")

    def finish(self, value):
        """
        Store user's choice and close dialog.

        Args:
            value: Button text that was clicked
        """
        self.result = value
        self.destroy()
        self.quit()


class CustomBox(Toplevel):
    """
    Custom message box with draggable interface.
    Used for displaying messages with custom button options.
    """

    def __init__(self, buttons, message, image_y):
        """
        Initialize custom dialog box.

        Args:
            buttons: List of button labels
            message: Message text to display
            image_y: Y-coordinate for image placement
        """
        super().__init__()
        self.result = None

        # Window configuration
        self.geometry("480x350")
        self.title("Mail Merge")
        self.resizable(True, True)
        self.attributes("-topmost", True)
        self.overrideredirect(True)  # Frameless window

        # Display label image
        self.photo = PhotoImage(file="label.png")
        image_label = Label(self, image=self.photo)
        image_label.place(x=128, y=image_y)

        # Bottom button bar
        bottom_frame = ttk.Frame(self)
        bottom_frame.place(relx=0, rely=1, anchor="sw", relwidth=1)

        # Create buttons
        for text in buttons:
            btn1 = ttk.Button(bottom_frame, text=text, command=lambda v=text: self.finish(v))
            btn1.pack(side="left", expand=True, fill="x", padx=5, pady=5)

        # Message text
        text = ttk.Label(self, text=message, font=("Segoe UI", 10))
        text.pack(side="top", pady=15)

        # Center window
        self.place_window_center()

        # Enable dragging
        self.bind("<Button-1>", self.start_move)
        self.bind("<B1-Motion>", self.do_move)

        self.mainloop()

    def start_move(self, event):
        """Record initial mouse position for dragging."""
        self._x = event.x
        self._y = event.y

    def do_move(self, event):
        """Calculate and apply window movement."""
        dx = event.x - self._x
        dy = event.y - self._y
        x = self.winfo_x() + dx
        y = self.winfo_y() + dy
        self.geometry(f"+{x}+{y}")

    def finish(self, value):
        """Store result and close dialog."""
        self.result = value
        self.destroy()
        self.quit()


class IntBox(Toplevel):
    """
    Custom dialog for integer input using a spinbox.
    Used for getting numeric values from user (e.g., recipient count).
    """

    def __init__(self, buttons, message):
        """
        Initialize integer input dialog.

        Args:
            buttons: List of button labels
            message: Prompt message for user
        """
        super().__init__()

        self.result = None  # Stores the integer value
        self.choice = ""  # Stores which button was clicked

        # Spinbox for number input (2-1000 range)
        self.spinbox = ttk.Spinbox(
            self,
            from_=2,
            to=1000,
            width=10
        )
        self.spinbox.pack(side="bottom", pady=60)

        # Window configuration
        self.geometry("580x400")
        self.title("Mail Merge")
        self.resizable(True, True)
        self.attributes("-topmost", True)
        self.overrideredirect(True)

        # Display image
        self.photo = PhotoImage(file="label.png")
        image_label = Label(self, image=self.photo)
        image_label.place(x=168, y=40)

        # Bottom button bar
        bottom_frame = ttk.Frame(self)
        bottom_frame.place(relx=0, rely=1, anchor="sw", relwidth=1)

        for text in buttons:
            btn1 = ttk.Button(bottom_frame, text=text, command=lambda v=text: self.finish(v))
            btn1.pack(side="left", expand=True, fill="x", padx=5, pady=5)

        # Message label
        text = ttk.Label(self, text=message, font=("Segoe UI", 10))
        text.place(relx=0.5, y=20, anchor='n')

        self.place_window_center()

        # Enable dragging
        self.bind("<Button-1>", self.start_move)
        self.bind("<B1-Motion>", self.do_move)

        self.mainloop()

    def start_move(self, event):
        """Record initial position for dragging."""
        self._x = event.x
        self._y = event.y

    def do_move(self, event):
        """Move window based on drag."""
        dx = event.x - self._x
        dy = event.y - self._y
        x = self.winfo_x() + dx
        y = self.winfo_y() + dy
        self.geometry(f"+{x}+{y}")

    def finish(self, value):
        """Store spinbox value and button choice, then close."""
        self.result = int(self.spinbox.get())
        self.choice = value
        self.destroy()
        self.quit()


class StringBox(Toplevel):
    """
    Custom dialog for string input using an entry widget.
    Used for getting text input from user (e.g., recipient names).
    """

    def __init__(self, buttons, message):
        """
        Initialize string input dialog.

        Args:
            buttons: List of button labels
            message: Prompt message for user
        """
        super().__init__()

        self.result = None
        self.choice = ""

        # Entry widget for text input
        self.entry = ttk.Entry(self)
        self.entry.pack(side="bottom", pady=60)

        # Window configuration
        self.geometry("580x400")
        self.title("Mail Merge")
        self.resizable(True, True)
        self.attributes("-topmost", True)
        self.overrideredirect(True)

        # Display image
        self.photo = PhotoImage(file="label.png")
        image_label = Label(self, image=self.photo)
        image_label.place(x=168, y=40)

        # Bottom button bar
        bottom_frame = ttk.Frame(self)
        bottom_frame.place(relx=0, rely=1, anchor="sw", relwidth=1)

        for text in buttons:
            btn1 = ttk.Button(bottom_frame, text=text, command=lambda v=text: self.finish(v))
            btn1.pack(side="left", expand=True, fill="x", padx=5, pady=5)

        # Message label
        text = ttk.Label(self, text=message, font=("Segoe UI", 10))
        text.place(relx=0.5, y=20, anchor='n')

        self.place_window_center()

        # Enable dragging
        self.bind("<Button-1>", self.start_move)
        self.bind("<B1-Motion>", self.do_move)

        self.mainloop()

    def start_move(self, event):
        """Record initial position for dragging."""
        self._x = event.x
        self._y = event.y

    def do_move(self, event):
        """Move window based on drag."""
        dx = event.x - self._x
        dy = event.y - self._y
        x = self.winfo_x() + dx
        y = self.winfo_y() + dy
        self.geometry(f"+{x}+{y}")

    def finish(self, value):
        """Store entry text and button choice, then close."""
        self.result = self.entry.get()
        self.choice = value
        self.destroy()
        self.quit()


class TextBox(Toplevel):
    """
    Custom dialog for multi-line text input.
    Used for entering letter content in Mail Merge.
    """

    def __init__(self, buttons, message):
        """
        Initialize text box dialog.

        Args:
            buttons: List of button labels
            message: Prompt message for user
        """
        super().__init__()

        self.result = None
        self.choice = ""

        # Multi-line text widget
        self.text_box = ttk.Text(self, width=300, height=100)
        self.text_box.pack(padx=50, pady=50)

        # Window configuration
        self.geometry("960x540")
        self.title("Mail Merge")
        self.resizable(True, True)
        self.attributes("-topmost", True)
        self.overrideredirect(True)

        # Bottom button bar
        bottom_frame = ttk.Frame(self)
        bottom_frame.place(relx=0, rely=1, anchor="sw", relwidth=1)

        for text in buttons:
            btn1 = ttk.Button(bottom_frame, text=text, command=lambda v=text: self.finish(v))
            btn1.pack(side="left", expand=True, fill="x", padx=5, pady=5)

        # Message label
        text = ttk.Label(self, text=message, font=("Segoe UI", 10))
        text.place(relx=0.5, y=20, anchor='n')

        self.place_window_center()

        # Enable dragging
        self.bind("<Button-1>", self.start_move)
        self.bind("<B1-Motion>", self.do_move)

        self.mainloop()

    def start_move(self, event):
        """Record initial position for dragging."""
        self._x = event.x
        self._y = event.y

    def do_move(self, event):
        """Move window based on drag."""
        dx = event.x - self._x
        dy = event.y - self._y
        x = self.winfo_x() + dx
        y = self.winfo_y() + dy
        self.geometry(f"+{x}+{y}")

    def finish(self, value):
        """Store text content and button choice, then close."""
        self.result = self.get_text()
        self.choice = value
        self.destroy()
        self.quit()

    def get_text(self):
        """
        Extract text content from text widget.

        Returns:
            String containing all text (without trailing newline)
        """
        content = self.text_box.get("1.0", "end-1c")  # '1.0' = start, 'end-1c' = exclude final newline
        return content


# =============================================================================
# LETTER CLASS
# =============================================================================

class Letter:
    """
    Handles mail merge letter generation.
    Creates personalized letters by replacing [name] placeholder with recipient names.
    """

    def __init__(self):
        """Initialize letter manager with empty state."""
        self.letter_body = ""  # Letter content template
        self.output_directory = ""  # Where to save generated letters
        self.letter_file_location = ""  # Source letter file path
        self.is_placeholder_not_present = False  # Validation flag

    def show_placeholder_instructions(self):
        """Display reminder about using [name] placeholder in letters."""
        Messagebox.show_info(
            title="Mail Merge",
            message="Make sure that you've entered the [name] placeholder in your letter"
        )

    def generate_personalized_letters(self):
        """
        Create individual letter files for each recipient.
        Replaces [name] placeholder with actual names and saves to disk.
        Handles duplicate filenames by appending (1).
        """
        # Prompt for output directory
        CustomBox(
            buttons=["Continue", "Exit"],
            message="Now please select the folders in which you want to save your mails.",
            image_y=40
        )

        self.output_directory = filedialog.askdirectory(
            title="Where do you want to save the mails?"
        )

        # Generate letter file for each recipient
        for recipient in name_manager.recipient_names:
            try:
                # Check if file already exists
                output_file = open(f"{self.output_directory}/{recipient}'s Mail.txt", "r")
                output_file.close()

                # File exists, create with (1) suffix
                with open(f"{self.output_directory}/{recipient}'s Mail(1).txt", "w") as output_file:
                    personalized_content = self.letter_body.replace("[name]", recipient)
                    output_file.write(personalized_content)

            except FileNotFoundError:
                # File doesn't exist, create normally
                with open(f"{self.output_directory}/{recipient}'s Mail.txt", "w") as output_file:
                    personalized_content = self.letter_body.replace("[name]", recipient)
                    output_file.write(personalized_content)

        # Show success message
        CustomBox(["Exit"], "Your Mail Merge was successful, congratulations!", 40)

        # Open output folder in system file explorer
        if platform.system() == "Windows":
            os.startfile(self.output_directory)
        elif platform.system() == "Darwin":  # macOS
            subprocess.call(["open", self.output_directory])
        else:  # Linux
            subprocess.call(["xdg-open", self.output_directory])

    def process_letter_content(self):
        """
        Handle letter content input (file browse or manual typing).
        Validates [name] placeholder presence before proceeding.
        Manages the workflow: input → validation → generation.
        """
        # Ask user how they want to provide letter content
        letter_choice_dialog = CustomBox(
            ["Browse", "Type Letter Content"],
            "How do you want to enter the letter content?\nAs a file or type it?",
            50
        )
        file_input_choice = letter_choice_dialog.result

        # Handle dialog cancellation
        if file_input_choice is None:
            show_exit_confirmation(self.process_letter_content)
            return

        # === BROWSE FILE OPTION ===
        if file_input_choice == "Browse":
            self.show_placeholder_instructions()

            # Get letter file from user
            self.letter_file_location = filedialog.askopenfilename(
                title="Select your mail",
                filetypes=(("Text files", "*.txt"),)
            )

            # Handle cancellation
            if not self.letter_file_location:
                show_exit_confirmation(self.process_letter_content)
                return

            # Read letter content
            with open(self.letter_file_location) as letter_file:
                self.letter_body = letter_file.read()

                # Validate [name] placeholder exists
                name_place = self.letter_body.find("[name]")
                if name_place < 0:
                    CustomBox(["Try Again"], "Warning: Placeholder '[name]' not present in file.", 40)
                    self.is_placeholder_not_present = True
                else:
                    self.is_placeholder_not_present = False

                # Re-prompt if placeholder missing
                if self.is_placeholder_not_present:
                    self.process_letter_content()
                    return

                # Proceed to generate letters
                self.generate_personalized_letters()

        # === TYPE CONTENT OPTION ===
        elif file_input_choice == "Type Letter Content":
            self.show_placeholder_instructions()

            # Get letter content via text box
            letter_body_dialog = TextBox(["Continue", "Exit"], "Enter the letter Content")
            self.letter_body = letter_body_dialog.result
            continue_choice = letter_body_dialog.choice

            # Handle exit choice
            if continue_choice == "Exit":
                show_exit_confirmation(self.process_letter_content)
                return

            # Check for empty letter
            elif continue_choice == "":
                Messagebox.show_warning(
                    title="Mail Merge",
                    message="The letter is empty, please try again."
                )
                self.process_letter_content()
                return

            # Validate [name] placeholder
            name_place = self.letter_body.find("[name]")
            if name_place < 0:
                Messagebox.show_warning(
                    title="Mail Merge",
                    message="Warning: Placeholder '[name]' is not in the letter body"
                )
                self.is_placeholder_not_present = True
            else:
                self.is_placeholder_not_present = False

            # Re-prompt if placeholder missing
            if self.is_placeholder_not_present:
                self.process_letter_content()
                return

            # Proceed to generate letters
            self.generate_personalized_letters()


# =============================================================================
# NAME MANAGER CLASS
# =============================================================================

class NameManager:
    """
    Manages recipient name collection for mail merge.
    Supports manual entry or text file import.
    Validates names contain only letters, spaces, and hyphens.
    """

    def __init__(self):
        """Initialize name manager with empty state."""
        self.recipient_names = []  # List of recipient names
        self.name_input_method = ""  # "Manual Insert" or "Text File"
        self.names_file_location = ""  # Path to names text file
        self.total_recipients = 0  # Number of recipients for manual entry

        # Prompt messages
        self.name_input_method_prompt = "Do you want to enter names manually, or through a text file?"
        self.recipient_count_prompt = "How many people do you want send a mail to?"
        self.name_prompt = "Please enter the name:"

    def name_collection(self):
        """
        Main workflow for name collection.
        Asks for input method, collects names, then confirms with user.
        """
        while True:
            # Ask user how they want to provide names
            name_input_method_dialog = CustomBox(
                buttons=["Manual Insert", "Text File", "Exit"],
                message=self.name_input_method_prompt,
                image_y=40
            )
            self.name_input_method = name_input_method_dialog.result

            # Handle cancellation
            if self.name_input_method is None:
                exit_confirmation()

            # Route to appropriate collection method
            if self.name_input_method == "Manual Insert":
                self.ask_recipient_count()
                self.collect_names_manually()

            elif self.name_input_method == "Text File":
                self.names_text_file_processing()

            else:  # Exit option
                exit_confirmation()

            # Show collected names for confirmation
            self.confirm_names_proceed()
            break

    def ask_recipient_count(self):
        """
        Prompt user for number of recipients (minimum 2).
        Validates input is at least 2 recipients.
        """
        while True:
            total_recipients_dialog_box = IntBox(["Continue", "Exit"], self.recipient_count_prompt)
            self.total_recipients = total_recipients_dialog_box.result

            # Handle exit choice
            if total_recipients_dialog_box.choice == "Exit":
                self.total_recipients = 0
                exit_confirmation()

            # Validate minimum count
            if self.total_recipients >= 2:
                break

    def collect_names_manually(self):
        """
        Collect names one-by-one through dialog boxes.
        Validates each name contains only: letters, spaces, hyphens.
        Converts names to title case (First Letter Capitalized).
        """
        # Collect remaining names (if some already exist)
        for _ in range(self.total_recipients - len(self.recipient_names)):
            while True:
                # Prompt for name
                entered_name_dialog = StringBox(["Continue", "Exit"], self.name_prompt)
                entered_name = entered_name_dialog.result
                name_dialog_exit_choice = entered_name_dialog.choice

                # Handle exit choice
                if name_dialog_exit_choice == "Exit":
                    show_exit_confirmation(self.collect_names_manually)
                    return

                # Validate name (only letters, spaces, hyphens allowed)
                if entered_name.replace(" ", "").replace("-", "").isalpha():
                    self.name_prompt = "Please enter the name:"  # Reset prompt
                    self.recipient_names.append(entered_name.title())  # Add in title case
                    break
                else:
                    # Invalid name, update prompt
                    self.name_prompt = "Please enter a valid name:"

    def ask_name_file(self):
        """
        Prompt user to select a text file containing names.
        Continues prompting until valid file is selected or user exits.
        """
        while True:
            self.names_file_location = ""
            self.names_file_location = filedialog.askopenfilename(
                title="Select the names text file",
                filetypes=(("Text files", "*.txt"),)
            )

            # Handle cancellation
            if not self.names_file_location:
                exit_confirmation()
            else:
                break

    def names_text_file_processing(self):
        """
        Read and parse names from text file.
        Expected format: "Name1, Name2, Name3, ..."
        Validates comma separation.
        """
        while True:
            self.ask_name_file()

            try:
                # Read file content
                with open(self.names_file_location) as names_file:
                    names_content = names_file.read()

            except Exception as e:
                # Handle file read errors
                Messagebox.show_warning(
                    title="Error occurred",
                    message=f"Error: {e}"
                )
                self.name_collection()
                return

            # Validate comma-separated format
            if ", " not in names_content:
                Messagebox.show_warning(
                    title="Mail Merge",
                    message="Please enter the names separated by comma."
                )
            elif ", " in names_content:
                break  # Valid format

        # Split names by comma and space
        self.recipient_names = names_content.split(", ")

    def confirm_names_proceed(self):
        """
        Display collected names to user for confirmation.
        Allows re-entry if names are incorrect, or proceed to letter entry.

        Returns:
            User's choice ("Continue", "Re-enter Names", or "Exit")
        """
        global confirmation_choice

        # Show names for confirmation
        names_confirmation_dialog = CustomBox(
            message=f"Names entered:\n{', '.join(self.recipient_names)}",
            buttons=["Continue", "Re-enter Names", "Exit"],
            image_y=60
        )

        confirmation_choice = names_confirmation_dialog.result

        # Handle user's choice
        if confirmation_choice == "Continue":
            return confirmation_choice

        elif confirmation_choice == "Re-enter Names":
            # Clear names and restart collection process
            self.recipient_names = []
            self.name_collection()
            return
        else:
            # Exit option selected
            exit_confirmation()


# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def center_window(window):
    """
    Center a window on the screen.

    Args:
        window: Tkinter window to center
    """
    window.update_idletasks()  # Ensure window dimensions are calculated
    width = window.winfo_width()
    height = window.winfo_height()
    x = (window.winfo_screenwidth() // 2) - (width // 2)
    y = (window.winfo_screenheight() // 2) - (height // 2)
    window.geometry(f'{width}x{height}+{x}+{y}')


def show_exit_confirmation(callback_function):
    """
    Show exit confirmation dialog.
    If user chooses not to exit, calls the provided callback function.

    Args:
        callback_function: Function to call if user chooses not to exit
    """
    exit_dialog = CustomBox(["Yes", "No"], "Do you want to exit the app?", 40)

    user_choice = exit_dialog.result

    if user_choice == "Yes":
        exit()  # Terminate application
    else:
        callback_function()  # Return to previous screen


def exit_confirmation():
    """
    Show exit confirmation dialog without callback.
    Simply exits if user confirms, otherwise returns to previous state.
    """
    exit_dialog = CustomBox(["Yes", "No"], "Do you want to exit the app?", 40)

    user_choice = exit_dialog.result

    if user_choice == "Yes":
        exit()
    else:
        pass  # Do nothing, return to caller


# =============================================================================
# INITIALIZATION & MAIN FLOW
# =============================================================================

# === INITIALIZE MAIN WINDOW ===
# Create hidden window (will be shown by text editor if selected)
app_window = ttk.Window(themename="darkly")
app_window.attributes('-alpha', 0)  # Start invisible
app_window.iconbitmap("logo.ico")
center_window(app_window)

# === INITIALIZE VARIABLES ===
confirmation_choice = ""  # Stores user's name confirmation choice
input_method_choice = ""  # Stores name input method choice

# === CREATE MANAGER OBJECTS ===
name_manager = NameManager()  # Handles recipient name collection
letter = Letter()  # Handles letter generation

# === MAIN APPLICATION SELECTION LOOP ===
while True:
    # Show starting screen for software selection
    start_screen = StartingWindow()
    software_choice = start_screen.result

    # Handle user's software choice
    if software_choice == "📧Mail Merge":
        break  # Proceed to Mail Merge
    elif software_choice == "📝Text Editor":
        break  # Proceed to Text Editor
    elif software_choice == "❌Exit":
        exit_confirmation()  # Confirm before exiting

# === LAUNCH SELECTED SOFTWARE ===

if software_choice == "📧Mail Merge":
    # === MAIL MERGE WORKFLOW ===
    # Step 1: Collect recipient names
    name_manager.name_collection()

    # Step 2: Process letter content and generate personalized letters
    letter.process_letter_content()

elif software_choice == "📝Text Editor":
    # === TEXT EDITOR WORKFLOW ===
    # Launch text editor application
    textEditor_launch()